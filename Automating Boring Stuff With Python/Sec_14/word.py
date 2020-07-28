import os
import docx
print(os.getcwd())
os.chdir('F:\\Python Projects\\Automating Boring Stuff With Python\\Sec_14')
print(os.getcwd())
d = docx.Document('demo.docx')
print(d.paragraphs[0].text)
p = d.paragraphs[1].text
# print(p.runs[1].text)  # it does not work somehow
# Just see the notes to check how to manipulate the doc. Google shit
